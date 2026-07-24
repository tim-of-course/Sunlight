from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "sunlight_open_alpha_feature_report.docx"


# Resolved preset: launch_messaging_guide (compact_reference_guide base).
# Named overrides: editorial_cover opening block and provisional Sunlight palette.
NAVY = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
AMBER = "C8841A"
PALE_AMBER = "FFF8E8"
PALE_BLUE = "E8EEF5"
PALE_GRAY = "F2F4F7"
INK = "1E2933"
MUTED = "5B6773"
GREEN = "2E6B4E"
RED = "9B1C1C"
WHITE = "FFFFFF"


def rgb(hex_value):
    return RGBColor.from_string(hex_value)


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge not in kwargs:
            continue
        edge_data = kwargs.get(edge)
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        for key in ["val", "sz", "space", "color"]:
            if key in edge_data:
                element.set(qn(f"w:{key}"), str(edge_data[key]))


def set_table_geometry(table, widths_dxa, indent_dxa=120, header_fill=None):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row_index, row in enumerate(table.rows):
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths_dxa[index] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_border(
                cell,
                top={"val": "single", "sz": "4", "color": "D7DEE7"},
                bottom={"val": "single", "sz": "4", "color": "D7DEE7"},
                left={"val": "single", "sz": "4", "color": "D7DEE7"},
                right={"val": "single", "sz": "4", "color": "D7DEE7"},
            )
            if row_index == 0 and header_fill:
                set_cell_shading(cell, header_fill)


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)
    set_run_font(run, size=9, color=MUTED)


def add_numbering(doc):
    numbering = doc.part.numbering_part.element
    abstract_id = 42
    num_id = 42
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    fmt = OxmlElement("w:numFmt")
    fmt.set(qn("w:val"), "bullet")
    text = OxmlElement("w:lvlText")
    text.set(qn("w:val"), "•")
    jc = OxmlElement("w:lvlJc")
    jc.set(qn("w:val"), "left")
    p_pr = OxmlElement("w:pPr")
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(ind)
    p_pr.append(tabs)
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Calibri")
    r_fonts.set(qn("w:hAnsi"), "Calibri")
    r_pr.append(r_fonts)
    level.extend([start, fmt, text, jc, p_pr, r_pr])
    abstract.append(level)
    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_num(paragraph, num_id, level=0):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), str(level))
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_el])


def set_keep(paragraph, keep_next=False, keep_lines=True):
    p_pr = paragraph._p.get_or_add_pPr()
    if keep_lines:
        p_pr.append(OxmlElement("w:keepLines"))
    if keep_next:
        p_pr.append(OxmlElement("w:keepNext"))


def add_para(doc, text="", style="Normal", before=None, after=None, align=None):
    paragraph = doc.add_paragraph(style=style)
    if align is not None:
        paragraph.alignment = align
    if before is not None:
        paragraph.paragraph_format.space_before = Pt(before)
    if after is not None:
        paragraph.paragraph_format.space_after = Pt(after)
    if text:
        run = paragraph.add_run(text)
        set_run_font(run, size=11, color=INK)
    return paragraph


def add_rich_para(doc, parts, style="Normal", before=None, after=None, align=None):
    paragraph = doc.add_paragraph(style=style)
    if align is not None:
        paragraph.alignment = align
    if before is not None:
        paragraph.paragraph_format.space_before = Pt(before)
    if after is not None:
        paragraph.paragraph_format.space_after = Pt(after)
    for text, options in parts:
        run = paragraph.add_run(text)
        set_run_font(run, size=options.get("size", 11), color=options.get("color", INK), bold=options.get("bold"), italic=options.get("italic"))
    return paragraph


def add_bullet(doc, text, num_id, level=0):
    paragraph = doc.add_paragraph(style="Normal")
    paragraph.paragraph_format.left_indent = Inches(0.375)
    paragraph.paragraph_format.first_line_indent = Inches(-0.188)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25
    apply_num(paragraph, num_id, level)
    run = paragraph.add_run(text)
    set_run_font(run, size=11, color=INK)
    return paragraph


def add_numbered_step(doc, label, text, num_id):
    paragraph = add_bullet(doc, text, num_id)
    # The numbered step is visually labeled in text while still using the real
    # numbering definition for a stable, editable list structure.
    paragraph._p.get_or_add_pPr().find(qn("w:numPr")).find(qn("w:numId")).set(qn("w:val"), str(num_id))
    return paragraph


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph(text, style=f"Heading {level}")
    set_keep(paragraph, keep_next=True)
    return paragraph


def add_kicker(doc, text):
    p = doc.add_paragraph(style="Kicker")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text.upper())
    set_run_font(run, size=10, color=AMBER, bold=True)
    return p


def set_paragraph_box(paragraph, fill, accent):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    for edge, size, color in (("top", "4", fill), ("left", "22", accent), ("bottom", "4", fill), ("right", "4", fill)):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "8")
        node.set(qn("w:color"), color)
        borders.append(node)
    p_pr.append(borders)


def add_callout(doc, label, text, fill=PALE_AMBER, accent=AMBER):
    p = doc.add_paragraph(style="Callout")
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(8)
    set_paragraph_box(p, fill, accent)
    p.style = "Callout"
    label_run = p.add_run(label + "  ")
    set_run_font(label_run, size=10.5, color=accent, bold=True)
    text_run = p.add_run(text)
    set_run_font(text_run, size=10.5, color=INK)
    return p


def add_table(doc, headers, rows, widths, header_fill=PALE_BLUE, font_size=9.5):
    table = doc.add_table(rows=1, cols=len(headers))
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        p = cell.paragraphs[0]
        p.style = "Table Text"
        run = p.add_run(header)
        set_run_font(run, size=font_size, color=NAVY, bold=True)
    for row_values in rows:
        row = table.add_row()
        for index, value in enumerate(row_values):
            cell = row.cells[index]
            p = cell.paragraphs[0]
            p.style = "Table Text"
            run = p.add_run(str(value))
            set_run_font(run, size=font_size, color=INK)
    set_table_geometry(table, widths, indent_dxa=120, header_fill=header_fill)
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)
    return table


def add_source_note(doc, text):
    p = doc.add_paragraph(style="Source Note")
    run = p.add_run(text)
    set_run_font(run, size=9, color=MUTED, italic=True)
    return p


def setup_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    title = styles["Title"]
    title.font.name = "Calibri"
    title._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    title._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    title.font.size = Pt(30)
    title.font.bold = True
    title.font.color.rgb = rgb(NAVY)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(8)
    title.paragraph_format.line_spacing = 1.0

    subtitle = styles["Subtitle"]
    subtitle.font.name = "Calibri"
    subtitle._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    subtitle._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    subtitle.font.size = Pt(15)
    subtitle.font.color.rgb = rgb(MUTED)
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(18)
    subtitle.paragraph_format.line_spacing = 1.15

    h1 = styles["Heading 1"]
    h1.font.name = "Calibri"
    h1._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    h1._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = rgb(BLUE)
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(10)
    h1.paragraph_format.line_spacing = 1.0

    h2 = styles["Heading 2"]
    h2.font.name = "Calibri"
    h2._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    h2._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.color.rgb = rgb(BLUE)
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(7)
    h2.paragraph_format.line_spacing = 1.0

    h3 = styles["Heading 3"]
    h3.font.name = "Calibri"
    h3._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    h3._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = rgb(DARK_BLUE)
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(5)
    h3.paragraph_format.line_spacing = 1.0

    for name in ["Kicker", "Callout", "Table Text", "Source Note", "Small Meta"]:
        if name not in styles:
            styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    styles["Kicker"].paragraph_format.space_before = Pt(0)
    styles["Kicker"].paragraph_format.space_after = Pt(16)
    styles["Callout"].paragraph_format.space_before = Pt(0)
    styles["Callout"].paragraph_format.space_after = Pt(0)
    styles["Callout"].paragraph_format.line_spacing = 1.15
    styles["Table Text"].paragraph_format.space_before = Pt(0)
    styles["Table Text"].paragraph_format.space_after = Pt(0)
    styles["Table Text"].paragraph_format.line_spacing = 1.10
    styles["Source Note"].paragraph_format.space_before = Pt(2)
    styles["Source Note"].paragraph_format.space_after = Pt(8)
    styles["Source Note"].paragraph_format.line_spacing = 1.1
    styles["Small Meta"].paragraph_format.space_before = Pt(0)
    styles["Small Meta"].paragraph_format.space_after = Pt(3)
    styles["Small Meta"].paragraph_format.line_spacing = 1.1


def setup_page(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("SUNLIGHT  /  OPEN ALPHA FEATURE REPORT")
    set_run_font(run, size=8.5, color=MUTED, bold=True)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run("Source-grounded working brief  |  Page ")
    set_run_font(run, size=9, color=MUTED)
    add_page_field(p)


def add_cover(doc):
    add_kicker(doc, "Open alpha feature report")
    title = doc.add_paragraph("Sunlight", style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph("A local-first source database for coding agents", style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_rich_para(
        doc,
        [("LANDING-PAGE SOURCE BRIEF", {"size": 10.5, "color": AMBER, "bold": True})],
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=28,
    )
    cover_meta = doc.add_paragraph(style="Small Meta")
    cover_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cover_meta.add_run("Repository-grounded draft  |  24 July 2026  |  Windows / NTFS alpha scope")
    set_run_font(r, size=10.5, color=MUTED)
    add_callout(
        doc,
        "READ FIRST",
        "The technical acceptance record says the Windows open alpha was approved, but the latest decision record says approval is suspended pending explicit-ignore source-completeness remediation. Treat this brief as messaging input, not a live release announcement, until that status is resolved and re-recorded.",
        fill=PALE_AMBER,
        accent=AMBER,
    )
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(40)
    spacer.paragraph_format.space_after = Pt(0)
    p = doc.add_paragraph(style="Small Meta")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Prepared for a downstream landing-page generation model")
    set_run_font(r, size=9.5, color=MUTED, italic=True)
    doc.add_page_break()


def build_report():
    doc = Document()
    setup_styles(doc)
    setup_page(doc)
    num_id = add_numbering(doc)
    add_cover(doc)

    add_heading(doc, "1. Executive summary", 1)
    add_para(doc, "Sunlight is a local-first source artifact system for coding agents. Its central idea is to replace ad hoc coordination through branches, worktrees, mutable directories, and guessed diffs with a native source database: agents inspect and mutate exact repository views through a small artifact API, each change belongs to a durable topic, topic revisions compose into immutable resolved views, tests run against those exact views, and accepted checkpoints can be exported as ordinary Git history.")
    add_para(doc, "The open-alpha story is strongest when framed as infrastructure for serious parallel agent work: fewer disposable checkouts, clearer ownership, deterministic composition, durable recovery, and an auditable path from an agent-authored change to a validated Git handoff. The product is intentionally local, technical, and opinionated. It is not a hosted collaboration suite, a GUI, a GitHub replacement, or a general-purpose secret scanner.")
    add_callout(doc, "ONE-SENTENCE THESIS", "Build MVCC for source code, but make the authoring surface artifact-native: isolated agents write durable topics through an API; exact topic combinations resolve into views; filesystem trees and Git commits are generated from those views when needed.", fill=PALE_BLUE, accent=BLUE)

    add_heading(doc, "2. Product snapshot", 1)
    add_table(
        doc,
        ["Field", "Source-grounded description"],
        [
            ("Product category", "Local-first source artifact system / native source-control database for coding agents."),
            ("Primary problem", "Parallel agent work creates worktree sprawl, branch/rebase ceremony, unclear context, and weak provenance."),
            ("Native source of truth", "Sunlight objects: artifacts, operations, topics, revisions, resolved views, executions, checkpoints, and evidence."),
            ("Primary interfaces", "CLI, repository-bound MCP server over local stdio, and a portable Agent Skill with Codex and Cursor adapters."),
            ("Current scope", "Local, single-repository workflows on Windows/NTFS for the open alpha."),
            ("Git relationship", "Git imports the baseline and receives approved checkpoint exports; Git is a compatibility surface, not native coordination."),
            ("Current release signal", "Acceptance evidence reports nine technical gates passed and a Windows approval, while the latest record suspends approval pending remediation review."),
        ],
        [2100, 7260],
    )
    add_source_note(doc, "Primary sources: README.md; docs/sunlight_consolidated_architecture_v0_3.md, sections 1-6; docs/open_alpha_acceptance.md, status and decision record; docs/acceptance/evidence/open_alpha_approval_2026-07-24.md and open_alpha_suspension_2026-07-24.md.")

    add_heading(doc, "3. The problem to make visible", 1)
    add_para(doc, "Coding agents are good at producing changes, but the default filesystem-and-Git workflow makes parallel work expensive to reason about. Each agent may need a separate checkout, branch, or worktree. The human or supervisor then has to reconstruct which context each change saw, whether two changes are compatible, whether a test ran against the exact combined result, and which generated outputs are safe to land.")
    add_para(doc, "Sunlight treats those coordination facts as product data rather than incidental developer memory. The result is a workflow where an agent can stay inside a precise context, own one intention, recover from ordinary failures, and hand off a validated result without making Git branches do the work of a database.")
    add_heading(doc, "What users should feel", 2)
    add_bullet(doc, "I can run several coding agents against the same logical repository without giving each one a full disposable checkout.", num_id)
    add_bullet(doc, "Each agent knows exactly which view it is reading and which topic owns its writes.", num_id)
    add_bullet(doc, "A stale hash or conflicting change becomes an inspectable fact with a recovery path, not a silent overwrite.", num_id)
    add_bullet(doc, "Tests and builds are tied to the exact resolved view they ran against.", num_id)
    add_bullet(doc, "When work is ready, I can freeze it and export a normal Git result that other tools understand.", num_id)

    add_heading(doc, "4. Feature inventory", 1)
    add_table(
        doc,
        ["Feature", "What it does", "Landing-page value", "Confidence / qualification"],
        [
            ("Native artifact authoring", "Read, list, search, patch, write, move, delete, metadata, and inspect through a topic-bound API.", "Agents can change source without treating a mutable project directory as the source of truth.", "Implemented and exercised in real-repository acceptance; say native authoring, not zero-filesystem."),
            ("Durable topics and revisions", "A topic names one coherent intention; each accepted operation advances an immutable topic revision.", "Parallel work stays legible and recoverable without branch ceremony.", "Core product model; avoid implying hosted collaboration."),
            ("Exact authoring sessions", "A session binds an actor to one write topic and one resolved view; accepted writes advance the session generation atomically.", "The agent sees its own latest work and can reason from a pinned context.", "Implemented; use exact view / pinned context language."),
            ("Deterministic view resolution", "Compose selected topic revisions into a conflict-free exact view or persistent conflict/staleness objects.", "Combine independent work deliberately and make overlap visible before execution.", "Implemented; do not promise automatic conflict repair."),
            ("Managed projections", "Materialize exact views for inspection, compatibility, execution, or export using measured strategies and cache reuse.", "Keep ordinary tools working without making projections the source of truth.", "Implemented; qualify Windows strategy and avoid blanket zero-copy claims."),
            ("Exact-view execution evidence", "Run a command against a resolved view and persist inputs, environment summary, outputs, timing, and result.", "Know what was tested, where, and against which tree identity.", "Implemented; network and some containment dimensions may be not enforced."),
            ("Explicit output promotion", "Convert approved source-like execution outputs into topic-owned operations with execution provenance.", "Generated or tool-produced changes become auditable instead of silently appearing in source.", "Implemented; promotion remains policy-gated."),
            ("Checkpoints and Git export", "Freeze a validated resolved view, run export policy, and create ordinary local Git history.", "Preserve familiar delivery while keeping native provenance intact.", "Implemented and acceptance-tested; no push or hosted forge in alpha."),
            ("Agent setup and doctor", "Install the portable skill and repository-bound Codex or Cursor MCP configuration; diagnose stale or missing setup.", "Reduce setup friction and give agents a safe next action when tools are unavailable.", "Implemented; generated paths are machine-local and clients require restart/reload."),
            ("Local MCP confinement", "Bind a stdio MCP server to one canonical repository root with typed tools, structured errors, bounded inputs, and durable IDs.", "Agents get a discoverable, repository-scoped workflow instead of arbitrary host access.", "Implemented; this is local stdio, not a network service or dashboard."),
        ],
        [1700, 2700, 2900, 2060],
        font_size=8.6,
    )
    add_source_note(doc, "Feature basis: docs/sunlight_consolidated_architecture_v0_3.md, sections 4-12; docs/local_mcp.md; integrations/agent-skills/sunlight/SKILL.md; crates/sun/src/mcp.rs tool schemas; crates/sunlight-core/src/resolver.rs, projection.rs, execution.rs, git_export.rs, and repo_state.rs.")

    add_heading(doc, "5. The end-to-end agent journey", 1)
    add_para(doc, "Use this as the narrative spine for a landing page. The user-facing story should stay plain-language first, with the CLI and MCP names available as proof for technical readers.")
    steps = [
        ("Install", "Run `sun agent install --client generic|codex|cursor` in the target repository. `sun agent doctor` checks the portable skill, client adapter, executable path, repository initialization, and restart requirement."),
        ("Initialize", "Run `sun init` or let the bound MCP server initialize an uninitialized repository. Sunlight imports the current Git baseline into native state and creates local state directories."),
        ("Name the work", "Create a durable topic such as `auth-fix` with an owner, visibility, and acceptance criteria. The topic is a user-facing intention, not a branch."),
        ("Open an exact session", "Start an actor-owned session over a resolved view. The session has one write topic, explicit capabilities, and a pinned frontier for predictable reads."),
        ("Inspect source", "Use artifact read, list, search, and inspect. Reads are scoped to the session or an exact resolved view and return hashes, IDs, classifications, and provenance."),
        ("Author with preconditions", "Patch or write through the artifact API using the exact content hash and path binding returned by Sunlight. A successful mutation creates one operation transaction and one new topic revision."),
        ("Resolve", "Select exact revisions from the base checkpoint. Independent changes compose deterministically; non-commutative overlap becomes a conflict or staleness object instead of being silently overwritten."),
        ("Run the real check", "Materialize an execution projection and run the command against the exact resolved view. Persist the view, tree identity, command, environment summary, bounded outputs, and result."),
        ("Freeze", "Create a checkpoint from the conflict-free resolved view and matching evidence. A checkpoint is immutable delivery evidence, not a moving branch pointer."),
        ("Hand off", "Validate export policy, then optionally create a normal local Git commit or branch. Native records remain the source of provenance; Git receives the compatibility result."),
    ]
    for label, text in steps:
        p = doc.add_paragraph(style="Normal")
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.space_after = Pt(6)
        lead = p.add_run(label + ". ")
        set_run_font(lead, size=11, color=NAVY, bold=True)
        body = p.add_run(text)
        set_run_font(body, size=11, color=INK)
    add_callout(doc, "SIMPLE FLOW", "Install -> initialize -> create topic -> start session -> read/search -> patch/write -> resolve exact view -> run tests -> checkpoint -> export to Git.", fill=PALE_BLUE, accent=BLUE)

    add_heading(doc, "6. Core concepts to explain without jargon", 1)
    add_table(
        doc,
        ["Term", "Plain-language explanation", "Why it matters"],
        [
            ("Artifact", "A source file or other durable item with a stable identity beyond its current path.", "Moves, deletes, and edits retain provenance instead of collapsing everything into filenames."),
            ("Operation transaction", "One accepted edit batch with exact before/after references and preconditions.", "Makes every change inspectable and safe to retry or reject."),
            ("Topic", "A durable change intention such as `auth-rework` or `profile-ui`.", "Gives concurrent agents a clear unit of ownership."),
            ("Resolved view", "One exact, immutable selection of a base plus specific topic revisions.", "Lets agents and tests talk about the same source state."),
            ("Projection", "A temporary ordinary-file representation generated from an exact view.", "Keeps existing tools usable without making the filesystem authoritative."),
            ("Execution", "A recorded command run against an exact view and its projection.", "Connects test evidence to the precise source tree that produced it."),
            ("Checkpoint", "A frozen validated view ready for review, landing, or export.", "Creates a durable handoff boundary."),
        ],
        [1700, 3700, 3960],
        font_size=9.1,
    )

    add_heading(doc, "7. Who this is for", 1)
    add_table(
        doc,
        ["Audience", "Current job-to-be-done", "Message that should land"],
        [
            ("Agentic software engineers", "Run multiple coding agents on one repository without worktree sprawl or hand-reconstructed context.", "Give each agent an exact world and a durable place to put its work."),
            ("Technical leads and supervisors", "Coordinate independent and overlapping changes, detect conflicts early, and validate the exact combined result.", "Parallelism is useful when the system can explain what happened."),
            ("Developer-tool builders", "Embed a repository-bound source API and structured lifecycle into coding clients.", "A small typed surface for authoring, resolution, execution, and handoff."),
            ("Infrastructure-minded teams on Windows", "Experiment with local-first agent workflows while retaining ordinary Git delivery.", "Adopt a new coordination layer without abandoning Git."),
            ("Researchers and evaluators", "Measure correctness, recovery, provenance, projection cost, and concurrent-agent behavior.", "The workflow produces durable evidence instead of demo-only state."),
        ],
        [1900, 3700, 3760],
        font_size=9.1,
    )
    add_heading(doc, "Not the primary audience yet", 2)
    add_bullet(doc, "Teams looking for a hosted forge, pull-request workflow, or remote collaboration product.", num_id)
    add_bullet(doc, "Users who need a polished graphical management interface or a background service.", num_id)
    add_bullet(doc, "Cross-platform teams that require macOS or Linux support in the initial alpha.", num_id)
    add_bullet(doc, "Organizations expecting Sunlight itself to replace dedicated secret scanning or repository hygiene controls.", num_id)

    add_heading(doc, "8. Positioning and differentiation", 1)
    add_para(doc, "The sharpest category distinction is not 'faster worktrees.' Sunlight is a source-control database with filesystem and Git adapters. The product changes the coordination substrate: native records are authoritative, filesystem projections are generated when a tool needs them, and Git is the familiar transport at the edge.")
    add_table(
        doc,
        ["Common workflow", "Sunlight framing"],
        [
            ("Each agent gets a branch/worktree and humans reconcile diffs later.", "Each agent gets a topic-bound session over an exact view; operations and revisions are durable as they happen."),
            ("A mutable checkout is assumed to be the truth.", "A resolved view is the truth; a projection is a generated adapter for tools."),
            ("A merge conflict is discovered after a sequence of Git operations.", "Non-commutative overlap is a first-class conflict/staleness object during view resolution."),
            ("A test result is attached loosely to a branch or local directory.", "Execution evidence is tied to the exact resolved view and tree identity."),
            ("Generated files appear as unexplained working-tree changes.", "Execution outputs are classified and require explicit promotion into a topic-owned operation."),
            ("Delivery means pushing a branch or asking another tool to reconstruct intent.", "A checkpoint freezes the handoff; optional Git export creates ordinary history with a persisted native-to-Git map."),
        ],
        [4200, 5160],
        font_size=9.1,
    )
    add_callout(doc, "POSITIONING RULE", "Lead with 'parallel agent work with exact context and durable provenance.' Explain the database model as the reason it works. Avoid leading with implementation terms such as MVCC, object stores, or content-addressed blobs unless the page has a technical deep-dive section.", fill=PALE_BLUE, accent=BLUE)

    add_heading(doc, "9. Trust, safety, and alpha guardrails", 1)
    add_para(doc, "Trust is a core part of the product story, but the page must be precise. Sunlight has strong boundaries around repository-relative paths, immutable records, exact hashes, conflict-free checkpoints, projection integrity, and Windows process containment. It does not provide universal sandboxing or secret prevention, and the current alpha status must be described as remediation-aware until the decision record is updated.")
    add_heading(doc, "What can be claimed", 2)
    add_bullet(doc, "The open-alpha target is Windows only, with Windows/NTFS as the tested scope.", num_id)
    add_bullet(doc, "The MCP server is bound to one canonical repository root and exposes typed, repository-relative operations.", num_id)
    add_bullet(doc, "Mutations use explicit preconditions and return structured errors with exact facts and next actions.", num_id)
    add_bullet(doc, "Conflict, staleness, projection, execution, checkpoint, and export state are persisted and inspectable.", num_id)
    add_bullet(doc, "Windows execution uses Job Object enforcement for the tested process/resource dimensions; policy reports disclose weaker dimensions.", num_id)
    add_bullet(doc, "The source inclusion contract follows normal Git semantics: tracked files are included, Git-ignored untracked files are excluded, and `.sunignore` can explicitly exclude additional paths.", num_id)
    add_heading(doc, "What must be qualified or avoided", 2)
    add_bullet(doc, "Do not say Sunlight scans, detects, blocks, or protects secrets. Secret prevention belongs to repository hygiene, permissions, deployment tooling, and dedicated scanners outside Sunlight.", num_id)
    add_bullet(doc, "Do not say macOS or Linux are supported in this alpha based on compilation or unit tests.", num_id)
    add_bullet(doc, "Do not say every projection is zero-copy, every run is fully isolated, or every policy dimension is enforced.", num_id)
    add_bullet(doc, "Do not say conflicts are auto-resolved. Sunlight makes conflicts inspectable; an agent or human still has to decide how to adapt the work.", num_id)
    add_bullet(doc, "Do not imply hosted collaboration, push automation, pull requests, a dashboard, cross-repository work, AST-native edits, or fuzzy rename inference are part of the alpha.", num_id)
    add_callout(doc, "CURRENT STATUS GUARDRAIL", "The repository contains a Windows approval record dated 2026-07-24 and a later suspension record from the same date. The landing page should use 'open alpha preparation' or 'Windows alpha, pending remediation review' until the owner confirms the release decision has been re-approved.", fill="FDECEC", accent=RED)

    add_heading(doc, "10. Evidence and proof points", 1)
    add_para(doc, "The technical evidence is strong enough to support a proof-oriented landing page, as long as the proof is scoped to the tested environment and not used to imply general availability or universal isolation.")
    add_table(
        doc,
        ["Proof point", "Evidence in repository", "Safe landing-page phrasing"],
        [
            ("Nine acceptance gates", "docs/open_alpha_acceptance.md and evidence bundle; checklist marks OA-01 through OA-09 complete for Windows/NTFS.", "Nine-gate acceptance journey completed for the declared Windows scope; release status remains subject to the latest remediation decision."),
            ("Realistic scale", "OA-07 final evidence: 5,818 tracked files, 90,935,221 logical bytes, four independent MCP authors, 20 mutations, four exact revisions.", "Validated with four concurrent authors on a 5,818-file repository."),
            ("Interactive latency", "OA-07 final evidence: status/list/search/read/mutation/resolution p95 373.752 / 125.717 / 1,092.236 / 123.886 / 1,317.980 / 155.161 ms.", "Measured interactive operations stayed within the frozen thresholds in the final Windows scale run."),
            ("Projection reuse", "OA-07 final evidence: first inspection 14,987 ms; cached verified inspection 1,081 ms / 7.21%; zero incremental physical bytes.", "Verified exact-view projections can be reused safely in the tested Windows path."),
            ("Validation", "OA-07 final evidence: two real passing `bun test ./src/lib` executions and two checkpoints; full source regression also passed.", "Tests and checkpoints were tied to exact resolved views, not mutable branches."),
            ("Fresh-client discovery", "OA-01 Cursor evidence plus fresh Codex/supervisor/unfamiliar-tester evidence.", "A fresh Cursor agent discovered the repository adapter from a natural engineering prompt and completed work without direct tracked-source access."),
            ("Git handoff", "OA-05 evidence and open_alpha_handoff.rs: exact checkpoint export to a safe, buildable local Git result.", "Accepted checkpoints can become ordinary Git history without requiring a push."),
        ],
        [1900, 3900, 3560],
        font_size=8.75,
    )
    add_source_note(doc, "Evidence sources: docs/acceptance/evidence/oa07_2026-07-23.json; oa07_attempts_2026-07-23.md; oa01_cursor_2026-07-24.md; oa01_oa03_oa09_fresh_clients_2026-07-24.md; oa02_oa06_2026-07-23.md; oa04_oa05_2026-07-23.md; crates/sun/tests/open_alpha_handoff.rs.")

    add_heading(doc, "11. Known limitations and post-alpha backlog", 1)
    add_table(
        doc,
        ["Limitation / deferral", "Implication for the page", "Suggested treatment"],
        [
            ("Windows only; Windows/NTFS tested scope", "This is a platform-specific alpha, not a cross-platform launch.", "Put it in the hero-adjacent scope note and FAQ."),
            ("Local, single repository", "No hosted workspace, multi-repo coordination, or remote service is included.", "Frame local-first as intentional; say single-repository explicitly."),
            ("No universal secret protection", "Tracked sensitive content is ordinary source unless excluded with `.sunignore` and external controls.", "Use a safety note; never use 'secure by default' as shorthand for secret handling."),
            ("Network may be not enforced in some real workflows", "A passing execution record can still report `network: not_enforced`.", "Show policy transparency as a feature; avoid 'offline sandbox' claims."),
            ("Compatibility projections are second-class", "Editing through projections requires explicit diff/import; projections are not source truth.", "Explain the native path first, then mention compatibility for legacy tools."),
            ("Absolute paths and client restarts", "Setup is local and may require restart/reload after MCP configuration changes.", "Include a concise install/doctor path and a troubleshooting FAQ."),
            ("No polished GUI or dashboard", "The primary experience is CLI + MCP + agent skill.", "Target technical early adopters and show commands or MCP cards."),
            ("Future work deferred", "Hosted forge, cross-repo intent trees, AST-native semantics, perfect dependency inference, fuzzy rename-plus-edit, and same-path multi-version filesystem remain future ideas.", "Keep in a roadmap section, not current feature claims."),
        ],
        [2200, 3500, 3660],
        font_size=8.9,
    )

    add_heading(doc, "12. Where this is going", 1)
    add_para(doc, "The current Windows alpha is the first focused product slice, not the final shape of Sunlight. The architecture and post-alpha notes point toward a broader coordination layer for agentic development. A landing page can show that arc, but it should label these as direction or roadmap rather than present them as available now.")
    add_callout(doc, "NOW / NEXT", "Now: local, single-repository, Windows/NTFS source coordination with exact views, durable provenance, execution evidence, checkpoints, and Git handoff. Next: extend the same exact-view model across repositories, richer tools, and more collaborative surfaces without giving up inspectability.", fill=PALE_BLUE, accent=BLUE)
    add_table(
        doc,
        ["Near-future direction", "What it could unlock", "How to phrase it today"],
        [
            ("Cross-repository intent trees", "Coordinate one coherent change across multiple repositories using the planned RepoTreeMap identity model.", "Designed to grow from one repository to multi-repository work; cross-repo support is not in the current alpha."),
            ("Richer compatibility and import flows", "Let legacy editors and agents work in projections with clearer diffs, explicit capture, and more reliable rename-plus-edit handling.", "Compatibility is an adapter today; the direction is a more capable bridge for tools that still need files."),
            ("A visual coordination surface", "Give humans a timeline or dashboard for topics, sessions, exact views, conflicts, executions, checkpoints, and exports.", "The model is designed for a clearer operational view; the current alpha is CLI/MCP-first and has no dashboard."),
            ("Broader client and editor integrations", "Bring the repository-bound workflow to more coding clients and editor experiences while preserving one portable skill and typed contracts.", "The portable skill and local MCP model are the integration foundation; current adapters are generic, Codex, and Cursor."),
            ("Semantic source operations", "Add AST- or symbol-aware edits and better dependency inference for changes that are difficult to express as file patches.", "The artifact model leaves room for structured entities and semantic operations; the alpha is file-oriented."),
            ("Smarter conflict adaptation", "Use explicit dependencies, richer conflict objects, and future assistive strategies to help agents adapt overlapping work without silent rewrites.", "Sunlight makes conflicts inspectable now; future versions can make resolution more assistive without making it opaque."),
            ("Live development and richer execution", "Support longer-lived dev servers, language servers, editor feedback, and more capable execution environments tied to exact views.", "The execution model is designed for live development later; the alpha focuses on bounded command runs and evidence."),
            ("Hosted forge and remote collaboration", "Connect accepted checkpoints and native provenance to review, sharing, and team-level delivery workflows.", "Git export is the local compatibility bridge today; hosted forge behavior is a post-alpha direction."),
        ],
        [2500, 3650, 3210],
        font_size=8.8,
        header_fill=PALE_AMBER,
    )
    add_heading(doc, "Roadmap language for the landing page", 2)
    add_bullet(doc, "Use a two-speed frame: 'Current Windows alpha' for what users can install and try now, and 'Where this is going' for the wider system being built.", num_id)
    add_bullet(doc, "Keep the through-line consistent: every future surface should make agent context, change ownership, composition, validation, and provenance more explicit.", num_id)
    add_bullet(doc, "Treat the roadmap as an architectural direction, not a promise of dates. Avoid dates, feature commitments, or platform claims that are not supplied by the owner.", num_id)
    add_bullet(doc, "Use visual contrast to separate current proof from future possibility: a solid 'Now' lane and a lighter 'Next' lane work well.", num_id)

    add_heading(doc, "13. Recommended landing-page architecture", 1)
    add_para(doc, "This is a content outline for the downstream model. The page can be redesigned after target audience and branding inputs arrive, but the proof sequence should remain recognizable.")
    page_sections = [
        ("Hero", "Promise: parallel coding agents need exact context, not more disposable checkouts. CTA: join / request access / follow the alpha, depending on the owner's availability decision."),
        ("Problem", "Show the cost of worktree sprawl, hidden context, merge archaeology, and test results that are hard to tie to one exact source state."),
        ("How it works", "Use the five-stage story: topic -> session -> exact view -> execution evidence -> checkpoint/Git handoff."),
        ("Feature proof", "Cards for native artifact authoring, durable topics, deterministic resolution, exact-view execution, output promotion, and Git export."),
        ("Agent workflow", "Show a compact command/MCP journey for technical credibility: init, topic, session, read, patch, resolve, run, checkpoint, export."),
        ("Trust and scope", "State Windows/NTFS, local single-repository scope, repository-bound MCP, structured errors, and the explicit source-inclusion contract."),
        ("Evidence", "Use the four-author / 5,818-file proof point, measured projection reuse, exact tests, and fresh-client acceptance evidence with a tested-scope label."),
        ("Roadmap / boundaries", "Use a clear Now / Next split: current Windows alpha capabilities first, then cross-repository intent trees, richer compatibility, a visual coordination surface, semantic operations, live development, and hosted review as future directions."),
        ("CTA and FAQ", "Answer: what is Sunlight, how does it relate to Git, what platforms are supported, does it replace secret scanning, and what does install require?"),
    ]
    for label, text in page_sections:
        p = doc.add_paragraph(style="Normal")
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(label + ". ")
        set_run_font(run, size=11, color=NAVY, bold=True)
        run = p.add_run(text)
        set_run_font(run, size=11, color=INK)

    add_heading(doc, "14. Provisional audience and brand guidance", 1)
    add_callout(doc, "OWNER INPUT", "Replace or refine this section with the audience, positioning, voice, visual identity, CTA, availability, and proof preferences you plan to provide to the landing-page model.", fill=PALE_AMBER, accent=AMBER)
    add_heading(doc, "Audience guidance", 2)
    add_bullet(doc, "Assume the reader is technically fluent but not already familiar with Sunlight's object model.", num_id)
    add_bullet(doc, "Speak first to people building or supervising agentic software workflows, not to generic productivity consumers.", num_id)
    add_bullet(doc, "Make the first conceptual leap small: 'exact context for every agent.' Introduce the database model only after the pain is clear.", num_id)
    add_bullet(doc, "Use a second layer for maintainers and infrastructure readers who want CLI, MCP, JSON, policy, and evidence details.", num_id)
    add_heading(doc, "Working voice", 2)
    add_bullet(doc, "Precise, calm, technical, and honest. Favor concrete verbs: bind, record, resolve, verify, freeze, export.", num_id)
    add_bullet(doc, "Confident without hype. Let durable evidence and visible limitations create credibility.", num_id)
    add_bullet(doc, "Avoid vague promises such as 'revolutionary collaboration,' 'frictionless AI coding,' or 'secure by default' unless the owner defines and substantiates them.", num_id)
    add_heading(doc, "Provisional visual direction", 2)
    add_table(
        doc,
        ["Element", "Working direction"],
        [
            ("Color idea", "Night navy (#0B2545) for trust and depth; warm sunlight amber (#C8841A) for accents; clear blue (#2E74B5) for interaction and technical proof; warm off-white backgrounds."),
            ("Typography", "Readable modern sans serif with strong hierarchy; use monospace only for short commands, IDs, and code fragments."),
            ("Imagery", "Abstract source trees, branching rays, layered exact views, and calm diagrams of topics resolving into a checkpoint. Prefer product schematics over generic robot imagery."),
            ("Motion", "If used, show a few source paths converging into one exact view, then a test result and Git handoff. Keep motion explanatory, not decorative."),
            ("Avoid", "Neon hacker aesthetics, cloud-control-plane clichés, Git logo mimicry, stock robot art, heavy matrix grids, and visuals that imply universal isolation or automatic conflict repair."),
        ],
        [1800, 7560],
        font_size=9.1,
        header_fill=PALE_AMBER,
    )

    add_heading(doc, "15. Owner inputs to add before publishing", 1)
    add_para(doc, "The downstream landing-page model should treat these as unresolved content fields rather than inventing answers:")
    for item in [
        "Primary audience and job title: who must feel the problem most urgently?",
        "Availability: is the alpha open to everyone, invite-only, waitlisted, or pending re-approval?",
        "Primary CTA: download, join waitlist, request access, install locally, or read the docs?",
        "Brand system: logo, final palette, typeface, imagery, tagline, and any existing site references.",
        "Preferred tone: research-grade, builder-focused, friendly, provocative, or enterprise-trustworthy.",
        "Proof assets: terminal captures, MCP tool cards, architecture diagrams, benchmark charts, or a short demo video.",
        "Commercial context: free/open source status, license language, sponsorship, or future pricing boundaries.",
        "Support expectations: where alpha users report issues, what machines are supported, and what the maintainer will respond to.",
    ]:
        add_bullet(doc, item, num_id)

    add_heading(doc, "16. Source basis and editorial notes", 1)
    add_para(doc, "This brief was assembled from the repository as of 24 July 2026. Facts below are the primary evidence set. Marketing interpretations are labeled by their role in the document; unresolved owner inputs are intentionally left as prompts.")
    add_table(
        doc,
        ["Source", "Used for"],
        [
            ("README.md", "Product definition, build/install flow, Windows-only scope, agent setup, Git relationship, and source inclusion contract."),
            ("docs/sunlight_consolidated_architecture_v0_3.md", "Product thesis, object model, lifecycle, projections, execution, Git interoperability, non-goals, and terminology."),
            ("docs/local_mcp.md", "MCP transport, tool list, repository binding, structured errors, lifecycle, request bounds, and current source-inclusion contract."),
            ("integrations/agent-skills/sunlight/SKILL.md and references", "Safe agent workflow, setup, recovery, and native authoring expectations."),
            ("docs/open_alpha_acceptance.md", "Acceptance gates, release criteria, platform scope, known limitations, and decision record; note current suspension status."),
            ("docs/acceptance/open_alpha_thresholds.md", "Frozen performance thresholds for the Windows/NTFS alpha."),
            ("docs/acceptance/evidence/*.md and *.json", "OA-01 through OA-09 evidence, fresh-client discovery, scale results, approval, and suspension history."),
            ("target/release/sun.exe --help", "Current CLI surface and the user-visible lifecycle vocabulary."),
            ("crates/sun/src/mcp.rs and lib.rs", "Typed MCP tool names, schemas, structured errors, CLI dispatch, and feature implementation paths."),
            ("crates/sunlight-core/src/{repo_state,resolver,projection,execution,git_export,compat_import}.rs", "Source-grounded implementation details for ingestion, exact views, projections, execution evidence, import/export, and policy boundaries."),
            ("crates/sun/tests/open_alpha_handoff.rs and self_hosting.rs", "Real-repository acceptance and exact Git handoff coverage."),
        ],
        [3600, 5760],
        font_size=8.8,
    )
    add_callout(doc, "EDITORIAL NOTE", "The repository is intentionally dirty with user changes that are part of the current remediation work. This report preserves that distinction: it summarizes the documented and implemented direction, but it does not treat uncommitted remediation as final release approval.", fill="FDECEC", accent=RED)

    # Keep the final source note compact and make the report easy to reuse as a handoff.
    p = doc.add_paragraph(style="Small Meta")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("End of report  |  Update the status guardrail and owner inputs before publication")
    set_run_font(r, size=9, color=MUTED, italic=True)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.core_properties.title = "Sunlight Open Alpha Feature Report"
    doc.core_properties.subject = "Landing-page source brief for Sunlight open alpha"
    doc.core_properties.author = "Sunlight project"
    doc.core_properties.keywords = "Sunlight, open alpha, coding agents, source artifacts, landing page"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_report()
